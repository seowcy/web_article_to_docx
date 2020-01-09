from goose3 import Goose # For Python 3
from bs4 import BeautifulSoup
import requests
import docx
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import sys


def get_data_from_urls(file_with_urls):
    with open(file_with_urls, 'r') as f:
        urls = [line.strip() for line in f.readlines() if line.strip()]
    data = {}
    counter = 1
    for url in urls:
        try:
            soup = BeautifulSoup(requests.get(url).content)
        except:
            article = "<ConnectionError>"
        if "thehackernews.com/" in url:
            title = soup.find("h1", {"class": "story-title"}).a.text.strip()
            article = soup.find("div", {"id": "articlebody"}).text.strip()
            article = re.sub(r"\(adsbygoogle.*\);", '', article)
            article = re.sub(r"\n[\n\s]*", '\n\n', article)
      
        data["TITLE%02d" % counter] = title
        data["ARTICLE%02d" % counter] = article
        data["SUMMARY%02d" % counter] = article.split('\n')[0]
        data["URL%02d" % counter] = url
        counter += 1
    return data, counter

def add_bookmark(run, bookmark_name):
    ''' Adds a word bookmark to a run '''
    tag = run._r
    start = docx.oxml.shared.OxmlElement('w:bookmarkStart')
    start.set(docx.oxml.ns.qn('w:id'), '0')
    start.set(docx.oxml.ns.qn('w:name'), bookmark_name)
    tag.append(start)

    text = docx.oxml.OxmlElement('w:r')
    tag.append(text)

    end = docx.oxml.shared.OxmlElement('w:bookmarkEnd')
    end.set(docx.oxml.ns.qn('w:id'), '0')
    end.set(docx.oxml.ns.qn('w:name'), bookmark_name)
    tag.append(end)

    return run

def main(file_with_urls, template, outfile):
    pattern = r"<.+>"
    mask = re.compile(pattern)
    document = Document(template)
    data, counter = get_data_from_urls(file_with_urls)

    t = document.tables[1]
    for i,row in enumerate(t.rows):
        for j,cell in enumerate(row.cells):
#             print(i,j,cell.text)
            match = mask.findall(cell.text)
            if match:
                if int(match[0][-3:-1]) < counter:
                    cell.text = data[match[0][1:-1]]

    flags = [0] * 10
    for i,p in enumerate(document.paragraphs):
#         print(i, p.text)
        match = mask.findall(p.text)
        if match:
            if int(match[0][-3:-1]) < counter:
                p.text = re.sub(pattern, data[match[0][1:-1]], p.text)
                p_format = p.paragraph_format
                p_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if match[0][1:-3] == "TITLE":
                p.add_run()
                if flags[int(match[0][-3:-1])-1] == 0:
                    add_bookmark(p.runs[0], "A%s" % int(match[0][-3:-1]))
                    p.runs[0].underline = True
                    flags[int(match[0][-3:-1])-1] = 1
                else:
                    add_bookmark(p.runs[0], "B%s" % int(match[0][-3:-1]))
                    p.runs[0].underline = True
                    p.runs[0].bold = True

    document.save(outfile)


if __name__ == '__main__':
    args = sys.argv[1:]
    if len(args) != 3:
        print("Usage: %s [file_with_urls] [template] [outfile]" % sys.argv[0])
        sys.exit()
    main(args[0], args[1], args[2])
